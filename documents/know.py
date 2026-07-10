#requirnments.py
Django>=5.1
djangorestframework

langchain
langchain-community
langchain-google-genai

google-generativeai

faiss-cpu

PyMuPDF
python-docx

python-dotenv

pillow
numpy

# Gemini API Key
from dotenv import load_dotenv
import os

load_dotenv()


# AI Settings
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200

TOP_K_RESULTS = 5

SUMMARY_MAX_CHARS = 30000

VECTOR_STORE_PATH = MEDIA_ROOT / "vector_store"


# document/constants.py
MAX_FILE_SIZE = 10 * 1024 * 1024

ALLOWED_FILE_TYPES = [
    "pdf",
    "doc",
    "docx",
]

SUMMARY_PROMPT = """
You are an expert document summarizer.

Generate a concise summary of the document.

Requirements:

- 3 to 5 paragraphs
- Use bullet points wherever helpful
- Preserve important information
- Keep names, dates and numbers accurate
- Do not add information that isn't in the document
"""

QUESTION_PROMPT = """
You are an AI assistant.

Answer ONLY using the provided context.

If the answer is not available inside the context,
reply with:

'I couldn't find that information in the document.'

Always mention page numbers whenever available.

Context:

{context}

Question:

{question}
"""


# document/enums.py
from django.db import models


class DocumentStatus(models.TextChoices):
    PROCESSING = "processing", "Processing"
    COMPLETED = "completed", "Completed"
    FAILED = "failed", "Failed"



# document/models.py
from django.db import models

from .enums import DocumentStatus


class Document(models.Model):
    file = models.FileField(upload_to="uploads/")

    file_name = models.CharField(max_length=255)

    file_type = models.CharField(max_length=10)

    extracted_content = models.TextField()

    summary = models.TextField(blank=True)

    status = models.CharField(
        max_length=20,
        choices=DocumentStatus.choices,
        default=DocumentStatus.PROCESSING,
    )

    uploaded_at = models.DateTimeField(auto_now_add=True)

    updated_at = models.DateTimeField(auto_now=True)

    class Meta:
        ordering = ["-uploaded_at"]

    def __str__(self):
        return self.file_name


class DocumentChunk(models.Model):
    document = models.ForeignKey(
        Document,
        on_delete=models.CASCADE,
        related_name="chunks",
    )

    chunk_index = models.PositiveIntegerField()

    page_number = models.PositiveIntegerField(
        null=True,
        blank=True,
    )

    chunk_text = models.TextField()

    class Meta:
        ordering = ["chunk_index"]

    def __str__(self):
        return f"{self.document.file_name} - Chunk {self.chunk_index}"
    


# chat/models.py
from django.db import models

from apps.document.models import Document


class ChatMessage(models.Model):
    document = models.ForeignKey(
        Document,
        on_delete=models.CASCADE,
        related_name="chat_messages",
    )

    question = models.TextField()

    answer = models.TextField()

    created_at = models.DateTimeField(auto_now_add=True)

    class Meta:
        ordering = ["created_at"]

    def __str__(self):
        return self.question[:50]
    

# document/admin.py
from django.contrib import admin

from .models import Document, DocumentChunk


@admin.register(Document)
class DocumentAdmin(admin.ModelAdmin):
    list_display = (
        "id",
        "file_name",
        "file_type",
        "status",
        "uploaded_at",
    )

    search_fields = (
        "file_name",
    )


@admin.register(DocumentChunk)
class DocumentChunkAdmin(admin.ModelAdmin):
    list_display = (
        "id",
        "document",
        "chunk_index",
        "page_number",
    )

# chat/admin.py
from django.contrib import admin

from .models import ChatMessage


@admin.register(ChatMessage)
class ChatMessageAdmin(admin.ModelAdmin):
    list_display = (
        "id",
        "document",
        "created_at",
    )

    search_fields = (
        "question",
    )


# #### part:2


document/
│
├── services.py          <- Upload workflow
├── ai_service.py        <- Gemini + LangChain
├── utils.py             <- PDF/DOC extraction
├── constants.py
├── enums.py


# document/utils.py
import fitz
from docx import Document as DocxDocument


def extract_pdf_text(file_path):
    """
    Returns:
    extracted_text,
    pages
    """

    document = fitz.open(file_path)

    full_text = ""
    pages = []

    for page_number, page in enumerate(document, start=1):
        text = page.get_text()

        full_text += text + "\n"

        pages.append(
            {
                "page": page_number,
                "text": text,
            }
        )

    return full_text.strip(), pages


def extract_docx_text(file_path):
    document = DocxDocument(file_path)

    text = "\n".join(
        paragraph.text
        for paragraph in document.paragraphs
    )

    return text.strip(), [
        {
            "page": 1,
            "text": text,
        }
    ]


def extract_document(file_path, file_type):
    file_type = file_type.lower()

    if file_type == "pdf":
        return extract_pdf_text(file_path)

    return extract_docx_text(file_path)


# document/ai_service.py

import os

from django.conf import settings

from langchain_google_genai import (
    ChatGoogleGenerativeAI,
    GoogleGenerativeAIEmbeddings,
)

from langchain.text_splitter import (
    RecursiveCharacterTextSplitter,
)

from langchain.docstore.document import Document

from langchain_community.vectorstores import FAISS

from .constants import (
    SUMMARY_PROMPT,
    QUESTION_PROMPT,
)



# Create LLM
llm = ChatGoogleGenerativeAI(
    model="gemini-2.5-flash",
    google_api_key=settings.GEMINI_API_KEY,
    temperature=0.2,
)


# Embedding Model
embedding_model = GoogleGenerativeAIEmbeddings(
    model="models/text-embedding-004",
    google_api_key=settings.GEMINI_API_KEY,
)


# Chunking
def split_document(document_pages):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.CHUNK_SIZE,
        chunk_overlap=settings.CHUNK_OVERLAP,
    )

    documents = []

    for page in document_pages:

        chunks = splitter.split_text(
            page["text"]
        )

        for chunk in chunks:

            documents.append(
                Document(
                    page_content=chunk,
                    metadata={
                        "page": page["page"],
                    },
                )
            )

    return documents



# Summary Generation

def generate_summary(text):

    if len(text) <= settings.SUMMARY_MAX_CHARS:
        return _single_summary(text)

    return _recursive_summary(text)


# Single Call
def _single_summary(text):

    prompt = f"""
    {SUMMARY_PROMPT}

    Document

    {text}
    """

    response = llm.invoke(prompt)

    return response.content

# Recursive Summary
def _recursive_summary(text):

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=12000,
        chunk_overlap=1000,
    )

    chunks = splitter.split_text(text)

    summaries = []

    for chunk in chunks:

        prompt = f"""
        {SUMMARY_PROMPT}

        {chunk}
        """

        response = llm.invoke(prompt)

        summaries.append(
            response.content
        )

    final_prompt = f"""
    Combine these summaries into
    one concise summary.

    {' '.join(summaries)}
    """

    final = llm.invoke(final_prompt)

    return final.content

# Create FAISS
def create_vector_store(
    document_id,
    chunks,
):

    vector_store = FAISS.from_documents(
        chunks,
        embedding_model,
    )

    save_path = os.path.join(
        settings.VECTOR_STORE_PATH,
        str(document_id),
    )

    vector_store.save_local(save_path)


# Load FAISS

def load_vector_store(document_id):

    path = os.path.join(
        settings.VECTOR_STORE_PATH,
        str(document_id),
    )

    return FAISS.load_local(
        path,
        embedding_model,
        allow_dangerous_deserialization=True,
    )



# Ask Question

def answer_question(
    document_id,
    question,
):

    vector_store = load_vector_store(
        document_id
    )

    docs = vector_store.similarity_search(
        question,
        k=settings.TOP_K_RESULTS,
    )

    context = ""

    for doc in docs:

        page = doc.metadata.get("page")

        context += (
            f"\n\nPage {page}\n"
            f"{doc.page_content}"
        )

    prompt = QUESTION_PROMPT.format(
        context=context,
        question=question,
    )

    response = llm.invoke(prompt)

    return response.content






# document/services.py
from .utils import extract_document

from .ai_service import (
    generate_summary,
    split_document,
    create_vector_store,
)

from .models import (
    DocumentChunk,
)

def process_document(document):

    extracted_text, pages = extract_document(
        document.file.path,
        document.file_type,
    )

    document.extracted_content = extracted_text

    summary = generate_summary(
        extracted_text
    )

    document.summary = summary

    chunks = split_document(
        pages
    )

    for index, chunk in enumerate(chunks):

        DocumentChunk.objects.create(
            document=document,
            chunk_index=index,
            page_number=chunk.metadata.get("page"),
            chunk_text=chunk.page_content,
        )

    create_vector_store(
        document.id,
        chunks,
    )

    document.status = "completed"

    document.save()


# final apps/document/utils.py

import fitz
from docx import Document as DocxDocument


class DocumentExtractor:
    """
    Extract text from supported document types.

    Returns:
        extracted_text: str
        pages: list[dict]

    Example pages:

    [
        {
            "page": 1,
            "text": "Page 1 content..."
        },
        {
            "page": 2,
            "text": "Page 2 content..."
        }
    ]
    """

    @staticmethod
    def extract(file_path: str, file_type: str):
        file_type = file_type.lower()

        if file_type == "pdf":
            return DocumentExtractor._extract_pdf(file_path)

        if file_type in ["doc", "docx"]:
            return DocumentExtractor._extract_docx(file_path)

        raise ValueError(f"Unsupported file type: {file_type}")

    @staticmethod
    def _extract_pdf(file_path: str):
        pdf = fitz.open(file_path)

        extracted_text = ""
        pages = []

        try:
            for page_number, page in enumerate(pdf, start=1):
                page_text = page.get_text("text").strip()

                extracted_text += page_text + "\n"

                pages.append(
                    {
                        "page": page_number,
                        "text": page_text,
                    }
                )

            return extracted_text.strip(), pages

        finally:
            pdf.close()

    @staticmethod
    def _extract_docx(file_path: str):
        document = DocxDocument(file_path)

        paragraphs = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()

            if text:
                paragraphs.append(text)

        extracted_text = "\n".join(paragraphs)

        pages = [
            {
                "page": 1,
                "text": extracted_text,
            }
        ]

        return extracted_text, pages
    


# apps/document/ai/prompts.py
"""
All AI prompts used by the application.

Keeping prompts in one place makes them easier to maintain
without touching the service classes.
"""

SUMMARY_PROMPT = """
You are an expert document summarizer.

Your task is to generate a concise and accurate summary of the
given document.

Instructions:

- Read the entire document carefully.
- Do NOT add information that is not present.
- Keep names, numbers and dates exactly as written.
- Focus only on the important information.
- Ignore repeated content.
- Keep the summary easy to read.

Return the summary in the following format:

## Overview

Write a short overview of the document in 2–3 paragraphs.

## Key Points

- Important point 1
- Important point 2
- Important point 3
- Continue as needed.

## Conclusion

Write a short concluding paragraph.

Document:

{document}
"""


CHUNK_SUMMARY_PROMPT = """
You are summarizing one small part of a larger document.

Create a concise summary of ONLY the provided text.

Do not mention that this is a chunk.

Text:

{document}
"""


FINAL_SUMMARY_PROMPT = """
Below are summaries generated from different parts of the same document.

Merge them into one final summary.

Requirements:

- Remove duplicate information.
- Keep all important facts.
- Maintain logical flow.
- Return the summary in the following format.

## Overview

2–3 paragraphs

## Key Points

- Point 1
- Point 2
- Point 3

## Conclusion

One short paragraph.

Summaries:

{summaries}
"""


QUESTION_PROMPT = """
You are an AI assistant that answers questions ONLY using the
provided document context.

Rules:

1. Answer only from the given context.
2. Never make up information.
3. If the answer cannot be found, reply exactly:

"I couldn't find this information in the document."

4. If page numbers are available, include them.
5. Keep the answer clear and concise.

Context:

{context}

Question:

{question}

Answer:
"""


#apps/document/ai/summary_service.py
from typing import List

from django.conf import settings

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import ChatGoogleGenerativeAI

from .prompts import (
    SUMMARY_PROMPT,
    CHUNK_SUMMARY_PROMPT,
    FINAL_SUMMARY_PROMPT,
)


class SummaryService:
    """
    Service responsible for generating document summaries.

    Strategy

    Small document
        -> One Gemini request

    Large document
        -> Split document
        -> Summarize every chunk
        -> Merge summaries
    """

    def __init__(self):
        self.llm = ChatGoogleGenerativeAI(
            model="gemini-2.5-flash",
            google_api_key=settings.GEMINI_API_KEY,
            temperature=0.2,
        )

        self.summary_threshold = settings.SUMMARY_MAX_CHARS

        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=12000,
            chunk_overlap=1000,
            separators=[
                "\n\n",
                "\n",
                ". ",
                " ",
                "",
            ],
        )

    def generate_summary(self, document_text: str) -> str:
        """
        Public method used by the document service.

        Automatically decides whether to generate the
        summary using a single request or recursive
        summarization.
        """

        document_text = document_text.strip()

        if not document_text:
            return ""

        if len(document_text) <= self.summary_threshold:
            return self._generate_single_summary(
                document_text
            )

        return self._generate_recursive_summary(
            document_text
        )

    def _generate_single_summary(
        self,
        document_text: str,
    ) -> str:
        """
        Generate summary using a single Gemini request.
        """

        prompt = SUMMARY_PROMPT.format(
            document=document_text
        )

        response = self.llm.invoke(prompt)

        if not response.content:
            return ""

        return response.content.strip()

    def _summarize_chunk(
        self,
        chunk: str,
    ) -> str:
        """
        Generate summary for a single chunk.

        Used during recursive summarization.
        """

        prompt = CHUNK_SUMMARY_PROMPT.format(
            document=chunk
        )

        response = self.llm.invoke(prompt)

        if not response.content:
            return ""

        return response.content.strip()
    
    def _generate_recursive_summary(
        self,
        document_text: str,
    ) -> str:
        """
        Generate summary for large documents.

        Workflow:

        Document
            ↓
        Split into chunks
            ↓
        Summarize each chunk
            ↓
        Merge chunk summaries
            ↓
        Return final summary
        """

        chunks = self.splitter.split_text(document_text)

        if not chunks:
            return ""

        chunk_summaries = []

        for chunk in chunks:
            summary = self._summarize_chunk(chunk)

            if summary:
                chunk_summaries.append(summary)

        if not chunk_summaries:
            return ""

        if len(chunk_summaries) == 1:
            return chunk_summaries[0]

        return self._merge_summaries(chunk_summaries)

    def _merge_summaries(
        self,
        summaries: List[str],
    ) -> str:
        """
        Merge all chunk summaries into one final summary.
        """

        prompt = FINAL_SUMMARY_PROMPT.format(
            summaries="\n\n".join(summaries)
        )

        response = self.llm.invoke(prompt)

        if not response.content:
            return ""

        return response.content.strip()
    

# apps/document/ai/vector_store.py
import os
from typing import List

from django.conf import settings

from langchain.docstore.document import Document
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_community.vectorstores import FAISS


class VectorStoreService:
    """
    Handles creation and loading of FAISS vector stores.
    """

    def __init__(self):
        self.embeddings = GoogleGenerativeAIEmbeddings(
            model="models/text-embedding-004",
            google_api_key=settings.GEMINI_API_KEY,
        )

    def split_document(self, pages: List[dict]) -> List[Document]:
        """
        Convert extracted pages into LangChain Documents.

        Each page is split into smaller chunks while preserving
        page number metadata.
        """

        from langchain.text_splitter import RecursiveCharacterTextSplitter

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
            separators=[
                "\n\n",
                "\n",
                ". ",
                " ",
                "",
            ],
        )

        documents = []

        for page in pages:
            chunks = splitter.split_text(page["text"])

            for chunk in chunks:
                documents.append(
                    Document(
                        page_content=chunk,
                        metadata={
                            "page": page["page"],
                        },
                    )
                )

        return documents

    def create_vector_store(
        self,
        document_id: int,
        chunks: List[Document],
    ) -> None:
        """
        Create FAISS index and save it to disk.
        """

        vector_store = FAISS.from_documents(
            documents=chunks,
            embedding=self.embeddings,
        )

        save_path = self._get_vector_path(document_id)

        os.makedirs(save_path, exist_ok=True)

        vector_store.save_local(save_path)

    def load_vector_store(
        self,
        document_id: int,
    ) -> FAISS:
        """
        Load FAISS vector store from disk.
        """

        save_path = self._get_vector_path(document_id)

        if not os.path.exists(save_path):
            raise FileNotFoundError(
                f"Vector store not found for document {document_id}"
            )

        return FAISS.load_local(
            folder_path=save_path,
            embeddings=self.embeddings,
            allow_dangerous_deserialization=True,
        )

    def similarity_search(
        self,
        document_id: int,
        query: str,
        top_k: int = None,
    ) -> List[Document]:
        """
        Retrieve the most relevant chunks for a question.
        """

        vector_store = self.load_vector_store(document_id)

        return vector_store.similarity_search(
            query=query,
            k=top_k or settings.TOP_K_RESULTS,
        )

    def _get_vector_path(
        self,
        document_id: int,
    ) -> str:
        """
        Returns:

        media/vector_store/12/
        """

        return os.path.join(
            settings.VECTOR_STORE_PATH,
            str(document_id),
        )
    
# apps/document/ai/qa_service.py

from typing import Dict, List

from django.conf import settings

from langchain.docstore.document import Document
from langchain_google_genai import ChatGoogleGenerativeAI

from .prompts import QUESTION_PROMPT
from .vector_store import VectorStoreService


class QAService:
    """
    Service responsible for answering questions from a document.

    Workflow:

        Question
            ↓
        Load FAISS
            ↓
        Retrieve Relevant Chunks
            ↓
        Build Context
            ↓
        Gemini
            ↓
        Answer + Sources
    """

    def __init__(self):
        self.llm = ChatGoogleGenerativeAI(
            model="gemini-2.5-flash",
            google_api_key=settings.GEMINI_API_KEY,
            temperature=0.2,
        )

        self.vector_service = VectorStoreService()

    def answer_question(
        self,
        document_id: int,
        question: str,
    ) -> Dict:
        """
        Returns

        {
            "answer": "...",
            "sources": [...]
        }
        """

        documents = self.vector_service.similarity_search(
            document_id=document_id,
            query=question,
        )

        context = self._build_context(documents)

        prompt = QUESTION_PROMPT.format(
            context=context,
            question=question,
        )

        response = self.llm.invoke(prompt)

        return {
            "answer": response.content.strip(),
            "sources": self._build_sources(documents),
        }

    def _build_context(
        self,
        documents: List[Document],
    ) -> str:
        """
        Convert retrieved chunks into prompt context.
        """

        context_parts = []

        for document in documents:

            page = document.metadata.get("page")

            context_parts.append(
                f"""
Page: {page}

Content:
{document.page_content}
"""
            )

        return "\n\n".join(context_parts)

    def _build_sources(
        self,
        documents: List[Document],
    ) -> List[dict]:
        """
        Prepare source metadata for API response.
        """

        sources = []

        seen_pages = set()

        for document in documents:

            page = document.metadata.get("page")

            if page in seen_pages:
                continue

            seen_pages.add(page)

            preview = document.page_content.strip()

            if len(preview) > 200:
                preview = preview[:200] + "..."

            sources.append(
                {
                    "page": page,
                    "preview": preview,
                }
            )

        return sources
    
#  apps/document/services.py
from django.db import transaction

from .models import Document, DocumentChunk
from .enums import DocumentStatus

from .utils import DocumentExtractor

from .ai.summary_service import SummaryService
from .ai.vector_store import VectorStoreService


class DocumentService:
    """
    Handles the complete document processing workflow.

    Workflow

        Upload
            ↓
        Extract Text
            ↓
        Generate Summary
            ↓
        Split Document
            ↓
        Save Chunks
            ↓
        Create FAISS
            ↓
        Update Status
    """

    def __init__(self):
        self.summary_service = SummaryService()
        self.vector_service = VectorStoreService()

    @transaction.atomic
    def process_document(
        self,
        document: Document,
    ) -> Document:
        """
        Process uploaded document.
        """

        try:
            extracted_text, pages = DocumentExtractor.extract(
                file_path=document.file.path,
                file_type=document.file_type,
            )

            document.extracted_content = extracted_text
            document.save(update_fields=["extracted_content"])

            summary = self.summary_service.generate_summary(
                extracted_text
            )

            document.summary = summary
            document.save(update_fields=["summary"])

            chunks = self.vector_service.split_document(
                pages
            )

            self._save_document_chunks(
                document,
                chunks,
            )

            self.vector_service.create_vector_store(
                document_id=document.id,
                chunks=chunks,
            )

            document.status = DocumentStatus.COMPLETED
            document.save(update_fields=["status"])

            return document

        except Exception:

            document.status = DocumentStatus.FAILED
            document.save(update_fields=["status"])

            raise

    def _save_document_chunks(
        self,
        document: Document,
        chunks,
    ):
        """
        Save all generated chunks into database.
        """

        DocumentChunk.objects.filter(
            document=document
        ).delete()

        chunk_objects = []

        for index, chunk in enumerate(chunks):

            chunk_objects.append(
                DocumentChunk(
                    document=document,
                    chunk_index=index,
                    page_number=chunk.metadata.get("page"),
                    chunk_text=chunk.page_content,
                )
            )

        DocumentChunk.objects.bulk_create(chunk_objects)


#  apps/document/serializers.py
import os

from rest_framework import serializers

from .constants import (
    ALLOWED_FILE_TYPES,
    MAX_FILE_SIZE,
)
from .enums import DocumentStatus
from .models import Document


class DocumentUploadSerializer(serializers.ModelSerializer):
    class Meta:
        model = Document
        fields = [
            "id",
            "file",
        ]
        read_only_fields = ["id"]

    def validate_file(self, file):
        extension = os.path.splitext(file.name)[1].lower().replace(".", "")

        if extension not in ALLOWED_FILE_TYPES:
            raise serializers.ValidationError(
                "Only PDF, DOC and DOCX files are allowed."
            )

        if file.size > MAX_FILE_SIZE:
            raise serializers.ValidationError(
                "File size must not exceed 10 MB."
            )

        return file

    def create(self, validated_data):
        uploaded_file = validated_data["file"]

        extension = os.path.splitext(
            uploaded_file.name
        )[1].lower().replace(".", "")

        return Document.objects.create(
            file=uploaded_file,
            file_name=uploaded_file.name,
            file_type=extension,
            status=DocumentStatus.PROCESSING,
        )


class DocumentListSerializer(serializers.ModelSerializer):
    class Meta:
        model = Document
        fields = [
            "id",
            "file_name",
            "file_type",
            "status",
            "uploaded_at",
        ]


class DocumentDetailSerializer(serializers.ModelSerializer):
    class Meta:
        model = Document
        fields = [
            "id",
            "file_name",
            "file_type",
            "summary",
            "status",
            "uploaded_at",
            "updated_at",
        ]


class DocumentChunkSerializer(serializers.Serializer):
    page = serializers.IntegerField()

    preview = serializers.CharField()


# apps/document/views.py
from rest_framework import status
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework.response import Response
from rest_framework.views import APIView

from .models import Document
from .serializers import (
    DocumentUploadSerializer,
    DocumentListSerializer,
    DocumentDetailSerializer,
)
from .services import DocumentService


class DocumentUploadAPIView(APIView):
    """
    Upload a document and process it.
    """

    parser_classes = (
        MultiPartParser,
        FormParser,
    )

    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.document_service = DocumentService()

    def post(self, request):
        serializer = DocumentUploadSerializer(
            data=request.data
        )

        serializer.is_valid(raise_exception=True)

        document = serializer.save()

        self.document_service.process_document(document)

        response_serializer = DocumentDetailSerializer(document)

        return Response(
            response_serializer.data,
            status=status.HTTP_201_CREATED,
        )


class DocumentListAPIView(APIView):
    """
    List all uploaded documents.
    """

    def get(self, request):
        documents = Document.objects.all()

        serializer = DocumentListSerializer(
            documents,
            many=True,
        )

        return Response(serializer.data)


class DocumentDetailAPIView(APIView):
    """
    Get a single document.
    """

    def get(self, request, pk):
        try:
            document = Document.objects.get(pk=pk)

        except Document.DoesNotExist:

            return Response(
                {
                    "detail": "Document not found."
                },
                status=status.HTTP_404_NOT_FOUND,
            )

        serializer = DocumentDetailSerializer(
            document
        )

        return Response(serializer.data)
    

# apps/document/urls.py
from django.urls import path

from .views import (
    DocumentUploadAPIView,
    DocumentListAPIView,
    DocumentDetailAPIView,
)

app_name = "document"

urlpatterns = [
    path(
        "upload/",
        DocumentUploadAPIView.as_view(),
        name="upload-document",
    ),
    path(
        "",
        DocumentListAPIView.as_view(),
        name="document-list",
    ),
    path(
        "<int:pk>/",
        DocumentDetailAPIView.as_view(),
        name="document-detail",
    ),
]


# apps/chat/models.py
from django.db import models

from apps.document.models import Document


class MessageType(models.TextChoices):
    USER = "user", "User"
    ASSISTANT = "assistant", "Assistant"


class ChatMessage(models.Model):
    """
    Stores question-answer pairs for a document.

    This model is also used as a cache. Before asking Gemini,
    we first check whether the same normalized question already
    exists for the document.
    """

    document = models.ForeignKey(
        Document,
        on_delete=models.CASCADE,
        related_name="chat_messages",
    )

    question = models.TextField()

    normalized_question = models.TextField(
        db_index=True,
    )

    answer = models.TextField()

    sources = models.JSONField(
        default=list,
        blank=True,
    )

    is_cached = models.BooleanField(
        default=False,
    )

    created_at = models.DateTimeField(
        auto_now_add=True,
    )

    updated_at = models.DateTimeField(
        auto_now=True,
    )

    class Meta:
        ordering = ["created_at"]

        indexes = [
            models.Index(
                fields=[
                    "document",
                    "normalized_question",
                ]
            )
        ]

    def __str__(self):
        return f"{self.document.file_name} - {self.question[:50]}"



# apps/chat/serializers.py
from rest_framework import serializers

from .models import ChatMessage


class AskQuestionSerializer(serializers.Serializer):
    """
    Serializer for asking a question about a document.
    """

    question = serializers.CharField(
        required=True,
        allow_blank=False,
        max_length=1000,
        trim_whitespace=True,
    )

    def validate_question(self, value):
        value = value.strip()

        if not value:
            raise serializers.ValidationError(
                "Question cannot be empty."
            )

        return value


class ChatMessageSerializer(serializers.ModelSerializer):
    """
    Serializer for chat history.
    """

    class Meta:
        model = ChatMessage
        fields = (
            "id",
            "question",
            "answer",
            "sources",
            "is_cached",
            "created_at",
        )


class AskQuestionResponseSerializer(serializers.Serializer):
    """
    Serializer for question response.
    """

    answer = serializers.CharField()

    sources = serializers.ListField(
        child=serializers.DictField(),
    )

    cached = serializers.BooleanField()


class ChatHistorySerializer(serializers.ModelSerializer):
    """
    Serializer for document chat history.
    """

    class Meta:
        model = ChatMessage
        fields = (
            "id",
            "question",
            "answer",
            "created_at",
        )


# apps/chat/services.py

import re
from typing import Dict, Optional

from django.db import transaction

from apps.document.models import Document
from apps.document.ai.qa_service import QAService

from .models import ChatMessage


class ChatService:
    """
    Handles all chat related business logic.

    Workflow

    User Question
          │
          ▼
    Normalize Question
          │
          ▼
    Check Cache
          │
      ┌───┴────┐
      │        │
    Found     Not Found
      │          │
      ▼          ▼
    Return     QA Service
                 │
                 ▼
           Save Chat
                 │
                 ▼
             Return Answer
    """

    def __init__(self):
        self.qa_service = QAService()

    @staticmethod
    def normalize_question(question: str) -> str:
        """
        Normalize user question so that similar
        questions can hit the cache.
        """

        question = question.lower()

        question = question.strip()

        question = re.sub(
            r"\s+",
            " ",
            question,
        )

        question = re.sub(
            r"[^\w\s]",
            "",
            question,
        )

        return question

    def get_document(
        self,
        document_id: int,
    ) -> Document:
        """
        Return document or raise exception.
        """

        return Document.objects.get(
            pk=document_id,
        )

    def get_cached_answer(
        self,
        document: Document,
        normalized_question: str,
    ) -> Optional[ChatMessage]:
        """
        Returns cached response if available.
        """

        return (
            ChatMessage.objects
            .filter(
                document=document,
                normalized_question=normalized_question,
            )
            .first()
        )

    @transaction.atomic
    def ask_question(
        self,
        document_id: int,
        question: str,
    ) -> Dict:
        """
        Main entry point for asking a question.

        Steps

        1. Validate document
        2. Normalize question
        3. Check cache
        4. If found return cached answer
        5. Otherwise call QA service
        """

        document = self.get_document(
            document_id
        )

        normalized_question = self.normalize_question(
            question
        )

        cached_chat = self.get_cached_answer(
            document,
            normalized_question,
        )

        if cached_chat:

            return {
                "answer": cached_chat.answer,
                "sources": cached_chat.sources,
                "cached": True,
            }

        #
        # Cache Miss
        #
        # Part 2 continues from here...
        #

                qa_response = self.qa_service.answer_question(
            document_id=document.id,
            question=question,
        )

        chat_message = self._save_chat_message(
            document=document,
            question=question,
            normalized_question=normalized_question,
            answer=qa_response["answer"],
            sources=qa_response["sources"],
        )

        return {
            "answer": chat_message.answer,
            "sources": chat_message.sources,
            "cached": False,
        }

    def _save_chat_message(
        self,
        document: Document,
        question: str,
        normalized_question: str,
        answer: str,
        sources: list,
    ) -> ChatMessage:
        """
        Save a newly generated question/answer pair.
        """

        return ChatMessage.objects.create(
            document=document,
            question=question,
            normalized_question=normalized_question,
            answer=answer,
            sources=sources,
            is_cached=False,
        )

    def get_chat_history(
        self,
        document_id: int,
    ):
        """
        Return all chat messages for a document.
        """

        return (
            ChatMessage.objects
            .filter(document_id=document_id)
            .order_by("created_at")
        )

    def delete_chat_history(
        self,
        document_id: int,
    ) -> int:
        """
        Delete all chat history for a document.

        Returns:
            Number of deleted records.
        """

        deleted_count, _ = (
            ChatMessage.objects
            .filter(document_id=document_id)
            .delete()
        )

        return deleted_count

    def has_chat_history(
        self,
        document_id: int,
    ) -> bool:
        """
        Check whether the document has any chat history.
        """

        return ChatMessage.objects.filter(
            document_id=document_id
        ).exists()
    


#apps/chat/views.py
from django.shortcuts import get_object_or_404

from rest_framework import status
from rest_framework.response import Response
from rest_framework.views import APIView

from apps.document.models import Document
from apps.document.enums import DocumentStatus

from .serializers import (
    AskQuestionSerializer,
    AskQuestionResponseSerializer,
    ChatHistorySerializer,
)
from .services import ChatService


chat_service = ChatService()


class AskQuestionAPIView(APIView):
    """
    Ask a question about a document.
    """

    def post(self, request, document_id):
        serializer = AskQuestionSerializer(
            data=request.data
        )

        serializer.is_valid(raise_exception=True)

        document = get_object_or_404(
            Document,
            pk=document_id,
        )

        if document.status != DocumentStatus.COMPLETED:
            return Response(
                {
                    "detail": "Document processing is not completed yet."
                },
                status=status.HTTP_400_BAD_REQUEST,
            )

        response = chat_service.ask_question(
            document_id=document.id,
            question=serializer.validated_data["question"],
        )

        response_serializer = AskQuestionResponseSerializer(
            data=response
        )

        response_serializer.is_valid(
            raise_exception=True
        )

        return Response(
            response_serializer.data,
            status=status.HTTP_200_OK,
        )


class ChatHistoryAPIView(APIView):
    """
    Return chat history of a document.
    """

    def get(self, request, document_id):
        document = get_object_or_404(
            Document,
            pk=document_id,
        )

        history = chat_service.get_chat_history(
            document.id
        )

        serializer = ChatHistorySerializer(
            history,
            many=True,
        )

        return Response(
            serializer.data,
            status=status.HTTP_200_OK,
        )


class ClearChatHistoryAPIView(APIView):
    """
    Delete all chat history for a document.
    """

    def delete(self, request, document_id):
        document = get_object_or_404(
            Document,
            pk=document_id,
        )

        deleted = chat_service.delete_chat_history(
            document.id
        )

        return Response(
            {
                "message": "Chat history deleted successfully.",
                "deleted_records": deleted,
            },
            status=status.HTTP_200_OK,
        )
    
#

